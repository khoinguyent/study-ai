#!/usr/bin/env python3
"""
Simple standalone script for DOCX parsing and chunking
This script demonstrates basic DOCX text extraction and chunking without external dependencies
"""

import asyncio
import sys
import os
from pathlib import Path
from typing import List, Dict, Any

def check_docx_dependency():
    """Check if python-docx is available"""
    try:
        from docx import Document
        return True
    except ImportError:
        return False

def extract_docx_text(file_path: Path) -> Dict[str, Any]:
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document
        
        # Load the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    tables_text.append(' | '.join(row_text))
        
        # Combine all text
        full_text = '\n\n'.join(paragraphs + tables_text)
        
        # Extract metadata
        metadata = {
            'paragraph_count': len(paragraphs),
            'table_count': len(doc.tables),
            'has_headers': any(paragraph.style.name.startswith('Heading') for paragraph in doc.paragraphs),
            'has_footers': len(doc.sections) > 0 and any(section.footer for section in doc.sections),
        }
        
        return {
            'success': True,
            'text': full_text,
            'metadata': metadata,
            'word_count': len(full_text.split()),
            'extraction_method': 'python-docx'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'text': '',
            'metadata': {},
            'word_count': 0
        }

def chunk_text(text: str, max_chunk_size: int = 1000) -> List[Dict[str, Any]]:
    """Chunk text into smaller pieces"""
    chunks = []
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    current_chunk = []
    current_size = 0
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        
        # If adding this paragraph would exceed chunk size, save current chunk
        if current_size + len(paragraph) > max_chunk_size and current_chunk:
            chunks.append({
                'content': '\n\n'.join(current_chunk),
                'size': current_size,
                'type': 'paragraph'
            })
            current_chunk = [paragraph]
            current_size = len(paragraph)
        else:
            current_chunk.append(paragraph)
            current_size += len(paragraph)
    
    # Add the last chunk if it has content
    if current_chunk:
        chunks.append({
            'content': '\n\n'.join(current_chunk),
            'size': current_size,
            'type': 'paragraph'
        })
    
    return chunks

async def test_docx_processing():
    """Test DOCX processing pipeline"""
    print("🚀 Simple DOCX Parsing and Chunking Test")
    print("=" * 50)
    
    # Check dependency
    if not check_docx_dependency():
        print("❌ python-docx library not available")
        print("   Install with: pip install python-docx")
        return
    
    print("✅ python-docx library available")
    
    # Path to the DOCX file
    docx_path = Path("data/Đồng Bằng Sông Cửu Long.docx")
    
    if not docx_path.exists():
        print(f"❌ DOCX file not found: {docx_path}")
        return
    
    print(f"📄 Processing: {docx_path}")
    print(f"📊 File size: {docx_path.stat().st_size / 1024:.1f} KB")
    print()
    
    try:
        # Step 1: Extract text
        print("🔍 Step 1: Text Extraction")
        print("-" * 25)
        
        extraction_result = extract_docx_text(docx_path)
        
        if not extraction_result['success']:
            print(f"❌ Extraction failed: {extraction_result.get('error')}")
            return
        
        print("✅ Text extraction successful!")
        print(f"📝 Word count: {extraction_result['word_count']}")
        print(f"🔧 Method: {extraction_result['extraction_method']}")
        
        # Show metadata
        metadata = extraction_result.get('metadata', {})
        if metadata:
            print("📊 Metadata:")
            for key, value in metadata.items():
                print(f"   {key}: {value}")
        
        print()
        
        # Step 2: Show text preview
        print("📖 Step 2: Text Preview")
        print("-" * 25)
        
        text = extraction_result['text']
        print(f"Text length: {len(text)} characters")
        print()
        
        # Show first 300 characters
        preview = text[:300]
        if len(text) > 300:
            preview += "..."
        print("Preview:")
        print("-" * 15)
        print(preview)
        print("-" * 15)
        print()
        
        # Step 3: Chunk the text
        print("✂️  Step 3: Text Chunking")
        print("-" * 25)
        
        chunks = chunk_text(text, max_chunk_size=1000)
        print(f"✅ Created {len(chunks)} chunks")
        print()
        
        # Step 4: Analyze chunks
        print("📊 Step 4: Chunk Analysis")
        print("-" * 25)
        
        total_chars = 0
        chunk_sizes = []
        
        for i, chunk in enumerate(chunks):
            size = chunk['size']
            chunk_sizes.append(size)
            total_chars += size
            
            print(f"Chunk {i+1}:")
            print(f"  Size: {size} characters")
            print(f"  Type: {chunk['type']}")
            print(f"  Preview: {chunk['content'][:80]}...")
            print()
        
        # Statistics
        print("📈 Statistics:")
        print(f"  Total chunks: {len(chunks)}")
        print(f"  Total characters: {total_chars}")
        print(f"  Average size: {total_chars / len(chunks):.1f} characters")
        print(f"  Largest: {max(chunk_sizes)} characters")
        print(f"  Smallest: {min(chunk_sizes)} characters")
        print()
        
        # Step 5: Quality check
        print("🔍 Step 5: Quality Check")
        print("-" * 25)
        
        # Verify text preservation
        reconstructed = "".join(chunk['content'] for chunk in chunks)
        preserved = len(reconstructed.strip()) == len(text.strip())
        print(f"Text preservation: {'✅ Complete' if preserved else '❌ Incomplete'}")
        
        if not preserved:
            print(f"Original: {len(text)} chars")
            print(f"Reconstructed: {len(reconstructed)} chars")
            print(f"Difference: {abs(len(text) - len(reconstructed))} chars")
        
        # Check chunk sizes
        oversized = [c for c in chunk_sizes if c > 1000]
        if oversized:
            print(f"⚠️  {len(oversized)} chunks exceed 1000 char limit")
        else:
            print("✅ All chunks within size limits")
        
        print()
        
        # Step 6: Show sample chunks
        print("📋 Step 6: Sample Chunks")
        print("-" * 25)
        
        for i in range(min(2, len(chunks))):
            chunk = chunks[i]
            print(f"Chunk {i+1} (Full):")
            print("-" * 30)
            print(chunk['content'])
            print("-" * 30)
            print()
        
        if len(chunks) > 2:
            print(f"... and {len(chunks) - 2} more chunks")
        
        print("🎉 Test completed successfully!")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    asyncio.run(test_docx_processing())
