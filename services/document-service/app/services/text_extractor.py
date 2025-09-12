import asyncio
import io
import os
from typing import Optional, Dict, Any
from pathlib import Path
import logging
import re

# Import text extraction libraries
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX files cannot be processed.")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF files cannot be processed.")

# Try to import the new robust PDF parser
try:
    import fitz  # PyMuPDF
    ROBUST_PDF_AVAILABLE = True
except ImportError:
    ROBUST_PDF_AVAILABLE = False
    logging.warning("PyMuPDF (fitz) not available. Robust PDF parsing disabled.")

# OCR dependencies (optional)
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR dependencies not available. OCR fallback disabled.")

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.warning("pandas not available. Excel files cannot be processed.")

logger = logging.getLogger(__name__)

def _clean_text_enhanced(text: str) -> str:
    """Enhanced text cleaning for better quality"""
    if not text:
        return ""
    
    # Dehyphenation and paragraph join
    text = re.sub(r"-\s*\n", "", text)           # join hyphenated line-breaks
    text = re.sub(r"[ \t]+\n", "\n", text)       # trim trailing spaces
    text = re.sub(r"\n{3,}", "\n\n", text)       # collapse huge gaps
    text = re.sub(r"\s+", " ", text)             # normalize whitespace
    
    # Remove common PDF artifacts
    text = re.sub(r'\x00', '', text)             # null characters
    text = re.sub(r'\x0c', '', text)             # form feed
    
    # Clean up common PDF text issues
    text = text.replace('ﬁ', 'fi')               # ligatures
    text = text.replace('ﬂ', 'fl')
    text = text.replace('ﬀ', 'ff')
    text = text.replace('ﬃ', 'ffi')
    text = text.replace('ﬄ', 'ffl')
    
    # Remove lines that are just punctuation or symbols
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if line.strip() and len(line.strip()) > 2:
            cleaned_lines.append(line.strip())
        elif line.strip() and line.strip().isalnum():
            cleaned_lines.append(line.strip())
    
    text = '\n'.join(cleaned_lines)
    
    # Final cleanup
    text = text.strip()
    text = re.sub(r'\n\n\n+', '\n\n', text)     # normalize paragraph breaks
    
    return text

class TextExtractor:
    """Service for extracting text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_pdf_text if PDF_AVAILABLE else None,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text if DOCX_AVAILABLE else None,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_excel_text if EXCEL_AVAILABLE else None,
            'text/plain': self._extract_plain_text,
            'application/msword': self._extract_doc_text,  # Legacy .doc files
            # Image formats with OCR support
            'image/png': self._extract_image_text if OCR_AVAILABLE else None,
            'image/jpeg': self._extract_image_text if OCR_AVAILABLE else None,
            'image/jpg': self._extract_image_text if OCR_AVAILABLE else None,
            'image/tiff': self._extract_image_text if OCR_AVAILABLE else None,
            'image/bmp': self._extract_image_text if OCR_AVAILABLE else None,
            'image/gif': self._extract_image_text if OCR_AVAILABLE else None,
        }
    
    async def extract_text(self, file_content: bytes, content_type: str, filename: str = None) -> Dict[str, Any]:
        """
        Extract text from document content
        
        Args:
            file_content: Raw file bytes
            content_type: MIME type of the file
            filename: Optional filename for better error messages
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            # Get the appropriate extraction method
            extractor = self.supported_formats.get(content_type)
            
            if not extractor:
                raise ValueError(f"Unsupported content type: {content_type}")
            
            # Run extraction in executor to avoid blocking
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None, 
                extractor, 
                file_content, 
                filename
            )
            
            return {
                'success': True,
                'text': result.get('text', ''),
                'metadata': result.get('metadata', {}),
                'word_count': result.get('word_count', 0),
                'extraction_method': result.get('method', 'unknown')
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from {filename or 'unknown file'}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {},
                'word_count': 0
            }
    
    def _extract_docx_text(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            # Create a BytesIO object from the file content
            doc_stream = io.BytesIO(file_content)
            
            # Load the document
            doc = DocxDocument(doc_stream)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Combine all text
            full_text = '\n\n'.join(paragraphs + tables_text)
            
            # Extract metadata
            metadata = {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'has_headers': any(paragraph.style.name.startswith('Heading') for paragraph in doc.paragraphs),
                'has_footers': len(doc.sections) > 0 and any(section.footer for section in doc.sections),
            }
            
            return {
                'text': full_text,
                'metadata': metadata,
                'word_count': len(full_text.split()),
                'method': 'python-docx'
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX file {filename}: {str(e)}")
            raise Exception(f"DOCX text extraction failed: {str(e)}")
    
    def _extract_doc_text(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """
        Extract text from legacy .doc files
        Note: This is a basic implementation. For better support, consider using antiword or similar tools.
        """
        try:
            # For now, we'll return a placeholder since .doc files require additional tools
            # In production, you might want to use antiword, catdoc, or similar tools
            
            # Basic text extraction attempt (this is limited)
            text_content = file_content.decode('utf-8', errors='ignore')
            
            # Try to extract readable text (this is very basic)
            readable_chars = ''.join(char for char in text_content if char.isprintable() or char.isspace())
            
            # Clean up the text
            cleaned_text = ' '.join(readable_chars.split())
            
            return {
                'text': cleaned_text,
                'metadata': {
                    'note': 'Basic .doc extraction - consider using antiword for better results',
                    'file_size': len(file_content)
                },
                'word_count': len(cleaned_text.split()),
                'method': 'basic-doc'
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOC file {filename}: {str(e)}")
            raise Exception(f"DOC text extraction failed: {str(e)}")
    
    def _extract_pdf_text(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract text from PDF files with enhanced text processing. Falls back to OCR if needed."""
        
        # Try robust PDF parser first if available
        if ROBUST_PDF_AVAILABLE:
            try:
                # Since we can't use await in a sync method, we'll use the sync version
                result = self._extract_pdf_text_robust_sync(file_content, filename)
                if result['success'] and result['text'].strip():
                    return result
                logger.info(f"Robust PDF parser returned insufficient text, falling back to PyPDF2")
            except Exception as e:
                logger.warning(f"Robust PDF parser failed, falling back to PyPDF2: {e}")
        
        # Fallback to PyPDF2
        return self._extract_pdf_text_pypdf2_sync(file_content, filename)
    
    def _extract_pdf_text_robust_sync(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract text using PyMuPDF (fitz) with OCR fallback - synchronous version"""
        try:
            import fitz
            
            # Open PDF from bytes
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            pages = []
            total_txt = []
            ocr_used = False
            total_words = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = ""
                method_used = "text"
                
                # Try structured text extraction first
                try:
                    blocks = page.get_text("blocks")
                    if blocks:
                        text_parts = []
                        for block in blocks:
                            if len(block) >= 5:  # block[4] contains text
                                text_parts.append(block[4])
                        page_text = "\n".join(text_parts).strip()
                        
                        if len(page_text) < 25:
                            page_text = page.get_text("text").strip()
                            method_used = "text_plain"
                    else:
                        page_text = page.get_text("text").strip()
                        method_used = "text_plain"
                except Exception as e:
                    logger.warning(f"Text extraction failed for page {page_num + 1}: {e}")
                    page_text = page.get_text("text").strip()
                    method_used = "text_plain"
                
                # OCR fallback if needed
                if len(page_text) < 25 and OCR_AVAILABLE:
                    try:
                        # Use PyMuPDF rasterization
                        zoom = 300 / 72.0  # 300 DPI
                        mat = fitz.Matrix(zoom, zoom)
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        
                        # OCR with Vietnamese + English support
                        ocr_text = pytesseract.image_to_string(
                            img, 
                            lang="vie+eng", 
                            config="--psm 6 --oem 3"
                        )
                        
                        if len(ocr_text.strip()) > len(page_text):
                            page_text = ocr_text.strip()
                            method_used = "ocr"
                            ocr_used = True
                            
                    except Exception as e:
                        logger.warning(f"OCR failed for page {page_num + 1}: {e}")
                
                # Clean the text
                cleaned_text = _clean_text_enhanced(page_text)
                
                if cleaned_text.strip():
                    pages.append({
                        'page_number': page_num + 1,
                        'word_count': len(cleaned_text.split()),
                        'character_count': len(cleaned_text),
                        'has_content': True,
                        'method': method_used
                    })
                    total_txt.append(f"--- Page {page_num + 1} ---\n{cleaned_text}")
                    total_words += len(cleaned_text.split())
                else:
                    pages.append({
                        'page_number': page_num + 1,
                        'word_count': 0,
                        'character_count': 0,
                        'has_content': False,
                        'note': 'Page contains no extractable text',
                        'method': method_used
                    })
            
            doc.close()
            
            # Build result
            full_text = "\n\n".join(total_txt)
            
            return {
                'success': True,
                'text': full_text,
                'metadata': {
                    'extraction_method': 'PyMuPDF (fitz)',
                    'pages_processed': len(pages),
                    'total_words': total_words,
                    'ocr_used': ocr_used,
                    'page_details': pages,
                    'text_quality': {
                        'quality': 'good' if total_words > 100 else 'fair',
                        'score': min(100, (total_words / 10) + 50),
                        'word_count': total_words,
                        'page_count': len(pages)
                    }
                },
                'word_count': total_words,
                'method': 'fitz-robust'
            }
            
        except Exception as e:
            logger.error(f"Robust PDF extraction failed: {str(e)}")
            raise Exception(f"Robust PDF extraction failed: {str(e)}")
    
    def _extract_pdf_text_pypdf2_sync(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract text from PDF files using PyPDF2 with enhanced text processing. Falls back to OCR if needed."""
        try:
            # Create a BytesIO object from the file content
            pdf_stream = io.BytesIO(file_content)
            
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            pages = []
            total_txt = []
            total_words = 0
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                # Clean the text
                cleaned_text = _clean_text_enhanced(page_text)
                
                if cleaned_text.strip():
                    pages.append({
                        'page_number': page_num + 1,
                        'word_count': len(cleaned_text.split()),
                        'character_count': len(cleaned_text),
                        'has_content': True,
                        'method': 'PyPDF2'
                    })
                    total_txt.append(f"--- Page {page_num + 1} ---\n{cleaned_text}")
                    total_words += len(cleaned_text.split())
                else:
                    pages.append({
                        'page_number': page_num + 1,
                        'word_count': 0,
                        'character_count': 0,
                        'has_content': False,
                        'note': 'Page contains no extractable text',
                        'method': 'PyPDF2'
                    })
            
            # Build result
            full_text = "\n\n".join(total_txt)
            
            return {
                'success': True,
                'text': full_text,
                'metadata': {
                    'extraction_method': 'PyPDF2',
                    'pages_processed': len(pages),
                    'total_words': total_words,
                    'page_details': pages,
                    'text_quality': {
                        'quality': 'good' if total_words > 100 else 'fair',
                        'score': min(100, (total_words / 10) + 50),
                        'word_count': total_words,
                        'page_count': len(pages)
                    }
                },
                'word_count': total_words,
                'method': 'pypdf2'
            }
            
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {str(e)}")
            raise Exception(f"PyPDF2 extraction failed: {str(e)}")
    
    def _clean_pdf_text(self, raw_text: str) -> str:
        """Clean and improve PDF extracted text quality"""
        if not raw_text:
            return ""
        
        # Remove excessive whitespace while preserving structure
        lines = raw_text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Clean individual line
            cleaned_line = self._clean_pdf_line(line)
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        
        # Join lines with proper spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Final cleanup
        cleaned_text = self._final_pdf_text_cleanup(cleaned_text)
        
        return cleaned_text
    
    def _clean_pdf_line(self, line: str) -> str:
        """Clean individual line from PDF text"""
        if not line:
            return ""
        
        # Remove excessive whitespace
        line = ' '.join(line.split())
        
        # Remove common PDF artifacts
        line = line.replace('\x00', '')  # Null characters
        line = line.replace('\x0c', '')  # Form feed
        
        # Clean up common PDF text issues
        line = line.replace('ﬁ', 'fi')  # Ligatures
        line = line.replace('ﬂ', 'fl')
        line = line.replace('ﬀ', 'ff')
        line = line.replace('ﬃ', 'ffi')
        line = line.replace('ﬄ', 'ffl')
        
        # Remove lines that are just punctuation or symbols
        if len(line.strip()) <= 2 and not line.strip().isalnum():
            return ""
        
        return line.strip()
    
    def _final_pdf_text_cleanup(self, text: str) -> str:
        """Final cleanup of PDF text"""
        if not text:
            return ""
        
        # Remove excessive blank lines
        lines = text.split('\n')
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            if line.strip():
                cleaned_lines.append(line)
                prev_empty = False
            elif not prev_empty:
                cleaned_lines.append(line)
                prev_empty = True
        
        # Join and clean up spacing
        result = '\n'.join(cleaned_lines)
        
        # Remove leading/trailing whitespace
        result = result.strip()
        
        # Normalize paragraph breaks
        result = result.replace('\n\n\n', '\n\n')
        
        return result
    
    def _assess_pdf_text_quality(self, text: str, page_metadata: list) -> dict:
        """Assess the quality of extracted PDF text"""
        if not text:
            return {'score': 0, 'issues': ['No text extracted'], 'quality': 'poor'}
        
        issues = []
        score = 100
        
        # Check text length
        if len(text) < 100:
            issues.append('Very short text - may be image-based PDF')
            score -= 30
        
        # Check for common PDF issues
        if '\x00' in text:
            issues.append('Contains null characters')
            score -= 10
        
        if text.count('\n') > len(text) * 0.1:  # Too many line breaks
            issues.append('Excessive line breaks - poor text structure')
            score -= 15
        
        # Check word density
        words = text.split()
        if words:
            avg_word_length = sum(len(word) for word in words) / len(words)
            if avg_word_length < 2:
                issues.append('Very short words - possible OCR artifacts')
                score -= 20
        
        # Check page coverage
        pages_with_content = len([p for p in page_metadata if p.get('has_content', False)])
        total_pages = len(page_metadata)
        if total_pages > 0:
            coverage = pages_with_content / total_pages
            if coverage < 0.5:
                issues.append(f'Low text coverage: {coverage:.1%} of pages have content')
                score -= 25
        
        # Determine quality level
        if score >= 80:
            quality = 'excellent'
        elif score >= 60:
            quality = 'good'
        elif score >= 40:
            quality = 'fair'
        else:
            quality = 'poor'
        
        return {
            'score': max(0, score),
            'issues': issues,
            'quality': quality,
            'text_length': len(text),
            'word_count': len(words),
            'page_coverage': pages_with_content / total_pages if total_pages > 0 else 0
        }

    def _extract_pdf_text_with_ocr(self, file_content: bytes) -> tuple[str, dict]:
        """Perform OCR on each PDF page and return combined text and metadata."""
        if not OCR_AVAILABLE:
            return "", {"available": False}
        try:
            images = convert_from_bytes(file_content, fmt='png')
            ocr_texts = []
            ocr_pages = []
            total_words = 0
            for idx, img in enumerate(images, start=1):
                # Optional: convert to grayscale for better OCR
                if img.mode != 'L':
                    img = img.convert('L')
                text = pytesseract.image_to_string(img)
                cleaned = self._final_pdf_text_cleanup(text)
                ocr_texts.append(f"--- OCR Page {idx} ---\n{cleaned}")
                words = len(cleaned.split())
                total_words += words
                ocr_pages.append({
                    'page_number': idx,
                    'word_count': words,
                    'character_count': len(cleaned),
                    'has_content': bool(cleaned.strip())
                })
            combined = '\n\n'.join(ocr_texts)
            return combined, {
                'pages': ocr_pages,
                'total_words': total_words,
                'engine': 'tesseract'
            }
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return "", {'error': str(e), 'engine': 'tesseract'}
    
    def _get_pdf_processing_recommendations(self, quality_assessment: dict) -> list:
        """Get recommendations for improving PDF text extraction"""
        recommendations = []
        
        if quality_assessment['quality'] == 'poor':
            recommendations.append({
                'priority': 'high',
                'action': 'OCR Processing',
                'description': 'Consider using OCR (Optical Character Recognition) for image-based PDFs',
                'tools': ['Tesseract', 'Adobe Acrobat Pro', 'Online OCR services']
            })
            
            if quality_assessment.get('page_coverage', 0) < 0.5:
                recommendations.append({
                    'priority': 'high',
                    'action': 'Page Analysis',
                    'description': 'Many pages contain no text - may be scanned documents or images',
                    'tools': ['PDF page inspection tools', 'Image analysis software']
                })
        
        if quality_assessment.get('issues'):
            for issue in quality_assessment['issues']:
                if 'null characters' in issue:
                    recommendations.append({
                        'priority': 'medium',
                        'action': 'Text Cleaning',
                        'description': 'Remove null characters and other artifacts',
                        'tools': ['Text preprocessing scripts', 'PDF repair tools']
                    })
                
                if 'excessive line breaks' in issue:
                    recommendations.append({
                        'priority': 'medium',
                        'action': 'Text Structure Analysis',
                        'description': 'Analyze and fix text structure issues',
                        'tools': ['Text analysis tools', 'PDF layout analysis']
                    })
        
        if quality_assessment['score'] < 60:
            recommendations.append({
                'priority': 'medium',
                'action': 'Alternative Extraction',
                'description': 'Try different PDF text extraction libraries',
                'tools': ['pdfplumber', 'pdfminer', 'PyMuPDF (fitz)']
            })
        
        return recommendations
    
    def extract_pdf_with_fallback(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """
        Extract text from PDF with fallback strategies for poor quality text
        
        This method attempts to extract text and provides recommendations
        for improving extraction quality when needed.
        """
        try:
            # First attempt: Standard text extraction
            extraction_result = self._extract_pdf_text(file_content, filename)
            
            # Assess text quality
            quality_assessment = extraction_result['metadata']['text_quality']
            
            # Add recommendations if quality is poor
            if quality_assessment['quality'] in ['poor', 'fair']:
                recommendations = self._get_pdf_processing_recommendations(quality_assessment)
                extraction_result['metadata']['recommendations'] = recommendations
                extraction_result['metadata']['needs_improvement'] = True
                
                # Add alternative extraction suggestions
                extraction_result['metadata']['alternative_methods'] = [
                    {
                        'method': 'OCR Processing',
                        'description': 'Use OCR for image-based PDFs',
                        'estimated_accuracy': '85-95%',
                        'requirements': 'OCR software or service'
                    },
                    {
                        'method': 'Layout Analysis',
                        'description': 'Analyze PDF layout for better text extraction',
                        'estimated_accuracy': '70-85%',
                        'requirements': 'Advanced PDF processing tools'
                    },
                    {
                        'method': 'Manual Review',
                        'description': 'Review and manually correct extracted text',
                        'estimated_accuracy': '95-100%',
                        'requirements': 'Human review time'
                    }
                ]
            else:
                extraction_result['metadata']['needs_improvement'] = False
                extraction_result['metadata']['recommendations'] = []
            
            # Mark success for unified contract
            extraction_result['success'] = True
            return extraction_result
            
        except Exception as e:
            logger.error(f"PDF extraction with fallback failed for {filename}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {
                    'extraction_method': 'failed',
                    'quality': 'failed',
                    'recommendations': [
                        {
                            'priority': 'high',
                            'action': 'Error Resolution',
                            'description': f'Fix extraction error: {str(e)}',
                            'tools': ['Error analysis', 'PDF validation']
                        }
                    ]
                },
                'word_count': 0
            }
    
    def _extract_excel_text(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract text from Excel files"""
        try:
            # Create a BytesIO object from the file content
            excel_stream = io.BytesIO(file_content)
            
            # Read Excel file
            excel_data = pd.read_excel(excel_stream, sheet_name=None)
            
            # Extract text from all sheets
            text_content = []
            for sheet_name, sheet_data in excel_data.items():
                if not sheet_data.empty:
                    text_content.append(f"--- Sheet: {sheet_name} ---")
                    
                    # Convert DataFrame to text
                    for index, row in sheet_data.iterrows():
                        row_text = ' | '.join(str(cell) for cell in row if pd.notna(cell))
                        if row_text.strip():
                            text_content.append(row_text)
                    
                    text_content.append("")  # Empty line between sheets
            
            full_text = '\n'.join(text_content)
            
            # Extract metadata
            metadata = {
                'sheet_count': len(excel_data),
                'sheet_names': list(excel_data.keys()),
                'total_rows': sum(len(sheet) for sheet in excel_data.values()),
                'file_size': len(file_content)
            }
            
            return {
                'text': full_text,
                'metadata': metadata,
                'word_count': len(full_text.split()),
                'method': 'pandas'
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from Excel file {filename}: {str(e)}")
            raise Exception(f"Excel text extraction failed: {str(e)}")
    
    def _extract_plain_text(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, use utf-8 with error handling
                text_content = file_content.decode('utf-8', errors='ignore')
            
            # Clean up the text
            cleaned_text = text_content.strip()
            
            # Extract metadata
            metadata = {
                'encoding': encoding if 'encoding' in locals() else 'unknown',
                'line_count': len(cleaned_text.splitlines()),
                'file_size': len(file_content)
            }
            
            return {
                'text': cleaned_text,
                'metadata': metadata,
                'word_count': len(cleaned_text.split()),
                'method': 'plain-text'
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from plain text file {filename}: {str(e)}")
            raise Exception(f"Plain text extraction failed: {str(e)}")
    
    def _extract_image_text(self, file_content: bytes, filename: str = None) -> Dict[str, Any]:
        """Extract text from image files using OCR with Vietnamese and English support."""
        if not OCR_AVAILABLE:
            return {
                'success': False, 
                'error': 'OCR dependencies not available. Please install pytesseract and tesseract-ocr.', 
                'text': '', 
                'metadata': {},
                'word_count': 0
            }
        
        try:
            # Convert bytes to a PIL Image object
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to grayscale for better OCR
            if image.mode != 'L':
                image = image.convert('L')
            
            # Perform OCR with Vietnamese + English support
            text = pytesseract.image_to_string(
                image,
                lang="vie+eng",  # Vietnamese + English
                config="--psm 6 --oem 3"  # Page segmentation mode 6, OCR Engine mode 3
            )
            
            # Clean up the text
            cleaned_text = _clean_text_enhanced(text)
            
            # Extract metadata
            metadata = {
                'extraction_method': 'OCR (Tesseract)',
                'file_size': len(file_content),
                'image_format': image.format,
                'image_mode': image.mode,
                'image_size': image.size,
                'ocr_engine': 'tesseract',
                'ocr_languages': 'vie+eng',
                'ocr_config': '--psm 6 --oem 3',
                'text_quality': {
                    'quality': 'good' if len(cleaned_text.split()) > 10 else 'fair',
                    'score': min(100, len(cleaned_text.split()) * 2),
                    'word_count': len(cleaned_text.split()),
                    'character_count': len(cleaned_text)
                }
            }
            
            return {
                'success': True,
                'text': cleaned_text,
                'metadata': metadata,
                'word_count': len(cleaned_text.split()),
                'method': 'ocr-image'
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from image file {filename}: {str(e)}")
            return {
                'success': False, 
                'error': f'Image OCR failed: {str(e)}', 
                'text': '', 
                'metadata': {},
                'word_count': 0
            }
    
    def get_supported_formats(self) -> list:
        """Get list of supported content types"""
        return [content_type for content_type, extractor in self.supported_formats.items() if extractor is not None]
    
    def is_format_supported(self, content_type: str) -> bool:
        """Check if a content type is supported"""
        return content_type in self.supported_formats and self.supported_formats[content_type] is not None
