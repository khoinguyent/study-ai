#!/usr/bin/env python3
"""
Advanced Document Analyzer
This script provides comprehensive analysis of documents including:
- Text extraction from multiple formats
- Advanced chunking strategies
- Content analysis and statistics
- Quality assessment
"""

import asyncio
import sys
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import json
from datetime import datetime

def check_dependencies():
    """Check if required dependencies are available"""
    dependencies = {
        'python-docx': False,
        'PyPDF2': False,
        'pandas': False
    }
    
    try:
        from docx import Document
        dependencies['python-docx'] = True
    except ImportError:
        pass
    
    try:
        import PyPDF2
        dependencies['PyPDF2'] = True
    except ImportError:
        pass
    
    try:
        import pandas as pd
        dependencies['pandas'] = True
    except ImportError:
        pass
    
    return dependencies

class DocumentAnalyzer:
    """Advanced document analyzer with multiple format support"""
    
    def __init__(self):
        self.dependencies = check_dependencies()
        self.supported_formats = self._get_supported_formats()
    
    def _get_supported_formats(self) -> Dict[str, bool]:
        """Get supported document formats based on available libraries"""
        return {
            'docx': self.dependencies['python-docx'],
            'pdf': self.dependencies['PyPDF2'],
            'xlsx': self.dependencies['pandas'],
            'txt': True  # Always supported
        }
    
    def analyze_document(self, file_path: Path) -> Dict[str, Any]:
        """Analyze a document and return comprehensive results"""
        if not file_path.exists():
            return {'error': f'File not found: {file_path}'}
        
        file_extension = file_path.suffix.lower()
        
        # Determine content type and extraction method
        if file_extension == '.docx' and self.dependencies['python-docx']:
            return self._analyze_docx(file_path)
        elif file_extension == '.pdf' and self.dependencies['PyPDF2']:
            return self._analyze_pdf(file_path)
        elif file_extension == '.xlsx' and self.dependencies['pandas']:
            return self._analyze_excel(file_path)
        elif file_extension == '.txt':
            return self._analyze_text(file_path)
        else:
            return {'error': f'Unsupported format: {file_extension}'}
    
    def _analyze_docx(self, file_path: Path) -> Dict[str, Any]:
        """Analyze DOCX document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Combine all text
            full_text = '\n\n'.join(paragraphs + tables_text)
            
            # Analyze document structure
            structure = {
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'has_headers': any(paragraph.style.name.startswith('Heading') for paragraph in doc.paragraphs),
                'has_footers': len(doc.sections) > 0 and any(section.footer for section in doc.sections),
            }
            
            return {
                'success': True,
                'format': 'docx',
                'text': full_text,
                'structure': structure,
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'extraction_method': 'python-docx'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'docx'
            }
    
    def _analyze_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Analyze PDF document"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text_parts = []
                page_count = len(pdf_reader.pages)
                
                for page_num in range(page_count):
                    page = pdf_reader.pages[page_num]
                    text_parts.append(page.extract_text())
                
                full_text = '\n\n'.join(text_parts)
                
                # Extract metadata
                metadata = {}
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                        'modification_date': pdf_reader.metadata.get('/ModDate', '')
                    }
                
                return {
                    'success': True,
                    'format': 'pdf',
                    'text': full_text,
                    'structure': {
                        'pages': page_count,
                        'metadata': metadata
                    },
                    'word_count': len(full_text.split()),
                    'character_count': len(full_text),
                    'extraction_method': 'PyPDF2'
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'pdf'
            }
    
    def _analyze_excel(self, file_path: Path) -> Dict[str, Any]:
        """Analyze Excel document"""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            text_parts = []
            total_rows = 0
            total_columns = 0
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheet_text = df.to_string(index=False, header=True)
                text_parts.append(f"Sheet: {sheet_name}\n{sheet_text}")
                total_rows += len(df)
                total_columns = max(total_columns, len(df.columns))
            
            full_text = '\n\n'.join(text_parts)
            
            return {
                'success': True,
                'format': 'xlsx',
                'text': full_text,
                'structure': {
                    'sheets': len(sheet_names),
                    'sheet_names': sheet_names,
                    'total_rows': total_rows,
                    'total_columns': total_columns
                },
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'extraction_method': 'pandas'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'xlsx'
            }
    
    def _analyze_text(self, file_path: Path) -> Dict[str, Any]:
        """Analyze plain text document"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                return {
                    'success': False,
                    'error': 'Could not decode text with any supported encoding',
                    'format': 'txt'
                }
            
            # Count lines
            lines = text.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]
            
            return {
                'success': True,
                'format': 'txt',
                'text': text,
                'structure': {
                    'total_lines': len(lines),
                    'non_empty_lines': len(non_empty_lines),
                    'encoding': encoding
                },
                'word_count': len(text.split()),
                'character_count': len(text),
                'extraction_method': 'native'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'txt'
            }
    
    def chunk_text_advanced(self, text: str, max_chunk_size: int = 1000, 
                           overlap: int = 100, strategy: str = 'paragraph') -> List[Dict[str, Any]]:
        """Advanced text chunking with multiple strategies"""
        chunks = []
        
        if strategy == 'paragraph':
            chunks = self._chunk_by_paragraphs(text, max_chunk_size)
        elif strategy == 'sentence':
            chunks = self._chunk_by_sentences(text, max_chunk_size, overlap)
        elif strategy == 'fixed':
            chunks = self._chunk_by_fixed_size(text, max_chunk_size, overlap)
        else:
            chunks = self._chunk_by_paragraphs(text, max_chunk_size)
        
        return chunks
    
    def _chunk_by_paragraphs(self, text: str, max_chunk_size: int) -> List[Dict[str, Any]]:
        """Chunk text by paragraphs"""
        chunks = []
        paragraphs = text.split('\n\n')
        current_chunk = []
        current_size = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            if current_size + len(paragraph) > max_chunk_size and current_chunk:
                chunks.append({
                    'content': '\n\n'.join(current_chunk),
                    'size': current_size,
                    'type': 'paragraph',
                    'strategy': 'paragraph'
                })
                current_chunk = [paragraph]
                current_size = len(paragraph)
            else:
                current_chunk.append(paragraph)
                current_size += len(paragraph)
        
        if current_chunk:
            chunks.append({
                'content': '\n\n'.join(current_chunk),
                'size': current_size,
                'type': 'paragraph',
                'strategy': 'paragraph'
            })
        
        return chunks
    
    def _chunk_by_sentences(self, text: str, max_chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """Chunk text by sentences with overlap"""
        # Simple sentence splitting (can be improved with NLTK)
        sentences = text.replace('\n', ' ').split('. ')
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence = sentence.strip() + '.'
            
            if current_size + len(sentence) > max_chunk_size and current_chunk:
                chunks.append({
                    'content': ' '.join(current_chunk),
                    'size': current_size,
                    'type': 'sentence',
                    'strategy': 'sentence'
                })
                
                # Add overlap
                overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) >= 2 else ' '.join(current_chunk)
                current_chunk = [overlap_text, sentence]
                current_size = len(overlap_text) + len(sentence)
            else:
                current_chunk.append(sentence)
                current_size += len(sentence)
        
        if current_chunk:
            chunks.append({
                'content': ' '.join(current_chunk),
                'size': current_size,
                'type': 'sentence',
                'strategy': 'sentence'
            })
        
        return chunks
    
    def _chunk_by_fixed_size(self, text: str, max_chunk_size: int, overlap: int) -> List[Dict[str, Any]]:
        """Chunk text by fixed size with overlap"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # Try to break at word boundary
            if end < len(text):
                # Find last space before max_chunk_size
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    'content': chunk_text,
                    'size': len(chunk_text),
                    'type': 'fixed',
                    'strategy': 'fixed'
                })
            
            # Move start position with overlap
            start = max(start + 1, end - overlap)
        
        return chunks
    
    def analyze_chunks(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze chunk quality and distribution"""
        if not chunks:
            return {'error': 'No chunks to analyze'}
        
        sizes = [chunk['size'] for chunk in chunks]
        total_chars = sum(sizes)
        
        # Reconstruct text to check preservation
        reconstructed = ''.join(chunk['content'] for chunk in chunks)
        
        analysis = {
            'total_chunks': len(chunks),
            'total_characters': total_chars,
            'average_chunk_size': total_chars / len(chunks),
            'largest_chunk': max(sizes),
            'smallest_chunk': min(sizes),
            'size_distribution': {
                'under_500': len([s for s in sizes if s < 500]),
                '500_1000': len([s for s in sizes if 500 <= s < 1000]),
                'over_1000': len([s for s in sizes if s >= 1000])
            },
            'chunk_types': {},
            'strategies': {}
        }
        
        # Count chunk types and strategies
        for chunk in chunks:
            chunk_type = chunk.get('type', 'unknown')
            strategy = chunk.get('strategy', 'unknown')
            
            analysis['chunk_types'][chunk_type] = analysis['chunk_types'].get(chunk_type, 0) + 1
            analysis['strategies'][strategy] = analysis['strategies'].get(strategy, 0) + 1
        
        return analysis

async def main():
    """Main function to demonstrate document analysis"""
    print("🚀 Advanced Document Analyzer")
    print("=" * 50)
    
    # Initialize analyzer
    analyzer = DocumentAnalyzer()
    
    # Show supported formats
    print("📋 Supported Formats:")
    for format_name, supported in analyzer.supported_formats.items():
        status = "✅" if supported else "❌"
        print(f"  {format_name.upper()}: {status}")
    print()
    
    # Analyze the DOCX file
    docx_path = Path("data/Đồng Bằng Sông Cửu Long.docx")
    
    if docx_path.exists():
        print(f"🔍 Analyzing: {docx_path.name}")
        print("-" * 30)
        
        # Extract and analyze
        result = analyzer.analyze_document(docx_path)
        
        if result.get('success'):
            print(f"✅ Extraction successful")
            print(f"📝 Words: {result['word_count']:,}")
            print(f"🔤 Characters: {result['character_count']:,}")
            print(f"🔧 Method: {result['extraction_method']}")
            
            # Show structure info
            structure = result.get('structure', {})
            if 'paragraphs' in structure:
                print(f"📄 Paragraphs: {structure['paragraphs']}")
            if 'tables' in structure:
                print(f"📊 Tables: {structure['tables']}")
            
            print()
            
            # Test different chunking strategies
            strategies = ['paragraph', 'sentence', 'fixed']
            
            for strategy in strategies:
                print(f"✂️  Testing {strategy} chunking strategy:")
                print("-" * 35)
                
                chunks = analyzer.chunk_text_advanced(
                    result['text'], 
                    max_chunk_size=1000, 
                    overlap=100, 
                    strategy=strategy
                )
                
                analysis = analyzer.analyze_chunks(chunks)
                
                print(f"  Chunks created: {analysis['total_chunks']}")
                print(f"  Average size: {analysis['average_chunk_size']:.1f} chars")
                print(f"  Size range: {analysis['smallest_chunk']} - {analysis['largest_chunk']} chars")
                
                # Show size distribution
                dist = analysis['size_distribution']
                print(f"  Distribution: <500: {dist['under_500']}, 500-1000: {dist['500_1000']}, >1000: {dist['over_1000']}")
                print()
            
            # Save results to JSON
            output_file = f"docx_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2, default=str)
            
            print(f"💾 Analysis results saved to: {output_file}")
            
        else:
            print(f"❌ Analysis failed: {result.get('error')}")
    
    else:
        print(f"❌ File not found: {docx_path}")
    
    print("\n🎉 Document analysis completed!")

if __name__ == "__main__":
    asyncio.run(main())
